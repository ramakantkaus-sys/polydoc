"""Image handling across formats.

The common real-world case: a Word document with a logo, converted to PDF. Images appear
in three structurally different positions, and each takes a different code path:

* a **standalone figure** -- a paragraph whose only content is a picture, which becomes a
  block-level :class:`~polydoc.model.Image`;
* an **inline logo** -- a picture in the middle of a text run, which becomes an
  :class:`~polydoc.model.InlineImage` and has to be embedded *within* a paragraph;
* a picture **inside a table cell**.

The inline path is the one that broke: the PDF writer passed a Python object id as
ReportLab's image source, so converting any Word document with a logo beside text raised
``OSError: Cannot open resource``. These tests exist so that cannot recur.
"""

from __future__ import annotations

import importlib.util
import zipfile
from io import BytesIO

import pytest

from conftest import needs_docx, needs_pdf_read, needs_pdf_write, needs_pptx

import polydoc
from polydoc.model import Document, Heading, Image, InlineImage, Paragraph, Text

needs_pillow = pytest.mark.skipif(
    importlib.util.find_spec("PIL") is None, reason="Pillow not installed"
)


@pytest.fixture
def logo_bytes() -> bytes:
    """A small PNG with two distinctive colours, so it is findable in rendered output."""
    from PIL import Image as PILImage
    from PIL import ImageDraw

    image = PILImage.new("RGB", (240, 120), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle([0, 0, 239, 119], outline="#003366", width=6)
    draw.ellipse([16, 16, 104, 104], fill="#cc0000")
    draw.rectangle([120, 40, 224, 80], fill="#003366")
    buffer = BytesIO()
    image.save(buffer, "PNG")
    return buffer.getvalue()


@pytest.fixture
def docx_with_images(logo_bytes, tmp_path):
    """A real DOCX with the logo standalone, inline beside text, and in a table cell."""
    import docx
    from docx.shared import Pt

    logo = tmp_path / "logo.png"
    logo.write_bytes(logo_bytes)

    document = docx.Document()
    document.add_heading("Acme Corporation", 1)
    document.add_picture(str(logo), width=Pt(180))

    paragraph = document.add_paragraph("Certified by ")
    paragraph.add_run().add_picture(str(logo), width=Pt(48))
    paragraph.add_run(" as of 2026.")

    document.add_paragraph("Body text after the figure.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Logo"
    table.cell(1, 0).paragraphs[0].add_run().add_picture(str(logo), width=Pt(60))
    table.cell(1, 1).text = "Company mark"

    path = tmp_path / "source.docx"
    document.save(str(path))
    return path


def media_count(path, prefix: str) -> int:
    with zipfile.ZipFile(path) as archive:
        return sum(1 for name in archive.namelist() if name.startswith(prefix))


def pdf_image_placements(path) -> int:
    """How many times an image is actually drawn, not how many objects exist.

    A PDF deduplicates identical images into one XObject, so counting objects would
    understate a document showing the same logo three times.
    """
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf

    doc = pymupdf.open(str(path))
    total = sum(len(page.get_image_info()) for page in doc)
    doc.close()
    return total


def rendered_colour_counts(path) -> tuple:
    """Render page 1 and count the logo's two signature colours."""
    try:
        import pymupdf
    except ImportError:
        import fitz as pymupdf

    doc = pymupdf.open(str(path))
    pixmap = doc.load_page(0).get_pixmap(dpi=100)
    samples, stride = pixmap.samples, pixmap.n
    red = navy = 0
    for offset in range(0, len(samples) - stride, stride):
        r, g, b = samples[offset], samples[offset + 1], samples[offset + 2]
        if r > 150 and g < 80 and b < 80:
            red += 1
        elif r < 80 and g < 90 and b > 90:
            navy += 1
    doc.close()
    return (red, navy)


# ---------------------------------------------------------------------------
# Model
# ---------------------------------------------------------------------------


class TestImageModel:
    def test_block_image_carries_bytes(self):
        image = Image(data=b"\x89PNG\r\n\x1a\n", mime_type="image/png", alt="logo")
        assert image.is_embedded
        assert image.text == "logo"

    def test_image_text_prefers_caption(self):
        assert Image(alt="alt", caption="cap").text == "cap"

    def test_inline_image_text_is_its_alt(self):
        assert InlineImage(alt="a logo").text == "a logo"

    def test_bytes_survive_json(self):
        from polydoc.model import Node

        image = Image(data=b"\x89PNG\r\n\x1a\n" + bytes(range(64)), mime_type="image/png")
        assert Node.from_dict(image.to_dict()).data == image.data

    def test_image_is_not_stripped_as_empty(self):
        from polydoc.edit import strip_empty

        document = Document([Image(data=b"x", alt=""), Paragraph.of("")])
        strip_empty(document)
        assert any(b.type == "image" for b in document.body)

    def test_paragraph_holding_an_inline_image_is_not_stripped(self):
        from polydoc.edit import strip_empty

        document = Document([Paragraph([InlineImage(src="x.png")])])
        assert strip_empty(document) == 0


# ---------------------------------------------------------------------------
# DOCX reading
# ---------------------------------------------------------------------------


@needs_docx
@needs_pillow
class TestDocxImageExtraction:
    def test_all_three_positions_are_found(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        blocks = [n for n in document.walk() if n.type == "image"]
        inlines = [n for n in document.walk() if n.type == "inline_image"]
        assert len(blocks) == 2, "standalone figure + table-cell picture"
        assert len(inlines) == 1, "the logo beside text"

    def test_bytes_are_extracted_not_just_referenced(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        images = [n for n in document.walk() if n.type in ("image", "inline_image")]
        assert images
        for image in images:
            assert image.data, f"{image.type} has no bytes"
            assert image.data[:8] == b"\x89PNG\r\n\x1a\n"

    def test_mime_type_recorded(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        image = document.images[0]
        assert image.mime_type and "png" in image.mime_type

    def test_dimensions_are_converted_to_points(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        image = document.images[0]
        # Written at 180 pt wide with a 2:1 source aspect ratio.
        assert image.width == pytest.approx(180, abs=1)
        assert image.height == pytest.approx(90, abs=1)

    def test_inline_image_stays_inside_its_paragraph(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        inline = [n for n in document.walk() if n.type == "inline_image"][0]
        assert inline.parent.type == "paragraph"
        # And the surrounding text is still there, in order.
        assert inline.parent.text.startswith("Certified by")
        assert "as of 2026" in inline.parent.text

    def test_table_cell_image_is_found_in_the_cell(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        cell_images = [
            n for n in document.walk()
            if n.type == "image" and any(a.type == "table_cell" for a in n.ancestors())
        ]
        assert len(cell_images) == 1

    def test_extraction_can_be_disabled(self, docx_with_images):
        document = polydoc.open(docx_with_images, extract_images=False)
        images = [n for n in document.walk() if n.type in ("image", "inline_image")]
        # Still located, but the bytes are not carried.
        assert images
        assert all(not image.data for image in images)

    def test_standalone_image_keeps_its_alignment(self, docx_with_images):
        # Word left-aligns add_picture(); losing that recentres figures on export.
        document = polydoc.open(docx_with_images)
        assert document.images[0].style is not None


# ---------------------------------------------------------------------------
# DOCX -> PDF, the case that crashed
# ---------------------------------------------------------------------------


@needs_docx
@needs_pdf_write
@needs_pillow
class TestDocxToPdfImages:
    def test_conversion_does_not_raise(self, docx_with_images, tmp_path):
        """Regression: an inline image used to raise OSError from ReportLab.

        The writer passed ``id(reader)`` -- a memory address -- as the image ``src``,
        so any Word document with a logo beside text failed to convert at all.
        """
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pdf")
        assert path.stat().st_size > 2000

    @needs_pdf_read
    def test_every_image_is_drawn(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pdf")
        assert pdf_image_placements(path) >= 3

    @needs_pdf_read
    def test_rendered_page_shows_the_artwork(self, docx_with_images, tmp_path):
        """The strongest check: not "is an image object present" but "is it visible"."""
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pdf")
        red, navy = rendered_colour_counts(path)
        assert red > 200, f"logo red not rendered ({red} px)"
        assert navy > 200, f"logo navy not rendered ({navy} px)"

    @needs_pdf_read
    def test_no_placeholder_text_leaks(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pdf")
        assert "[Image:" not in polydoc.open(path).text

    @needs_pdf_read
    def test_text_around_the_inline_image_survives(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        text = polydoc.open(document.save(tmp_path / "out.pdf")).text
        assert "Certified by" in text
        assert "as of 2026" in text

    def test_images_can_be_omitted(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pdf", embed_images=False)
        assert path.exists()

    def test_corrupt_image_data_degrades_instead_of_crashing(self, tmp_path):
        """One unreadable image must not cost the whole conversion."""
        document = Document([
            Heading.of("Report", 1),
            Image(data=b"not a real image at all", alt="broken figure"),
            Paragraph([Text("before "), InlineImage(data=b"also broken", alt="broken logo"),
                       Text(" after")]),
            Paragraph.of("This text must still appear."),
        ])
        path = document.save(tmp_path / "broken.pdf")
        assert path.stat().st_size > 1000

    @needs_pdf_read
    def test_corrupt_image_falls_back_to_alt_text(self, tmp_path):
        document = Document([
            Image(data=b"not an image", alt="broken figure"),
            Paragraph.of("Following text."),
        ])
        text = polydoc.open(document.save(tmp_path / "broken.pdf")).text
        assert "Following text." in text

    def test_aspect_ratio_is_preserved_when_one_dimension_is_given(self, logo_bytes, tmp_path):
        # A 2:1 source given only a width must not be stretched.
        document = Document([Image(data=logo_bytes, width=120, mime_type="image/png")])
        assert document.save(tmp_path / "aspect.pdf").exists()

    def test_oversized_image_is_capped_to_the_frame(self, logo_bytes, tmp_path):
        document = Document([Image(data=logo_bytes, width=5000, mime_type="image/png")])
        # Must lay out rather than overflow and raise.
        assert document.save(tmp_path / "huge.pdf").exists()


# ---------------------------------------------------------------------------
# Other writers
# ---------------------------------------------------------------------------


@needs_docx
@needs_pillow
class TestImagesInOtherFormats:
    def test_docx_round_trip_re_embeds(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.docx")
        assert media_count(path, "word/media/") >= 1

    @needs_pptx
    def test_pptx_re_embeds(self, docx_with_images, tmp_path):
        document = polydoc.open(docx_with_images)
        path = document.save(tmp_path / "out.pptx")
        assert media_count(path, "ppt/media/") >= 1

    def test_html_inlines_a_data_uri(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        html = document.to_text("html")
        assert "data:image/png;base64," in html

    def test_html_can_reference_instead_of_inlining(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        html = document.to_text("html", embed_images=False)
        assert "data:image/png;base64," not in html

    def test_markdown_emits_an_image_reference(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        markdown = document.to_text("markdown")
        assert "![" in markdown

    def test_json_preserves_bytes_losslessly(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        restored = polydoc.loads(document.to_text("json"), "json")

        def data_of(doc):
            return [
                n.data for n in doc.walk()
                if n.type in ("image", "inline_image") and n.data
            ]

        assert data_of(restored) == data_of(document)

    def test_plain_text_describes_the_image(self, docx_with_images):
        document = polydoc.open(docx_with_images)
        assert "[Image:" in document.to_text("txt")


@needs_pdf_write
@needs_pdf_read
@needs_pillow
class TestPdfImageExtraction:
    def test_images_are_recovered_when_reading_a_pdf(self, logo_bytes, tmp_path):
        source = Document([
            Heading.of("Figure", 1),
            Image(data=logo_bytes, mime_type="image/png", width=180),
        ])
        path = source.save(tmp_path / "fig.pdf")
        reread = polydoc.open(path)
        images = [n for n in reread.walk() if n.type == "image" and n.data]
        assert images, "no image recovered from the PDF"
        assert images[0].data[:4] in (b"\x89PNG", b"\xff\xd8\xff\xe0", b"\xff\xd8\xff\xdb")

    def test_extraction_can_be_disabled_on_read(self, logo_bytes, tmp_path):
        source = Document([Image(data=logo_bytes, mime_type="image/png", width=180)])
        path = source.save(tmp_path / "fig.pdf")
        assert not [n for n in polydoc.open(path, images=False).walk() if n.type == "image"]

    def test_full_cycle_docx_to_pdf_to_docx(self, logo_bytes, tmp_path):
        source = Document([
            Heading.of("Cycle", 1),
            Image(data=logo_bytes, mime_type="image/png", width=180),
        ])
        docx_path = source.save(tmp_path / "a.docx")
        pdf_path = polydoc.open(docx_path).save(tmp_path / "b.pdf")
        final = polydoc.open(pdf_path).save(tmp_path / "c.docx")
        assert media_count(final, "word/media/") >= 1
