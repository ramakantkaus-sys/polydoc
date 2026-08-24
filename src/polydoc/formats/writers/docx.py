"""DOCX writer, built on python-docx.

Two areas need XML-level work because python-docx has no high-level API for them:

**Hyperlinks.** ``add_run`` cannot produce a ``w:hyperlink``, so the relationship is
created directly and the element assembled by hand.

**Nested lists.** python-docx applies list styles by name ("List Bullet 2"), which only
reaches three levels and does not restart numbering. We write a real ``numPr`` with the
right ``ilvl``, backed by a numbering definition created on demand.

Everything else maps cleanly: headings use the built-in ``Heading N`` styles so Word's
navigation pane and generated tables of contents work.
"""

from __future__ import annotations

from typing import Any, BinaryIO, Dict, List, Optional, Sequence, Tuple

from ...model import (
    Alignment,
    Block,
    BlockContainer,
    CodeBlock,
    Container,
    Document,
    DynamicField,
    Footnote,
    FootnoteRef,
    Heading,
    HorizontalRule,
    Image,
    Inline,
    InlineImage,
    LineBreak,
    Link,
    ListBlock,
    ListItem,
    ListStyle,
    Math,
    Page,
    PageBreak,
    Paragraph,
    ParagraphStyle,
    Quote,
    Section,
    Slide,
    Table,
    TableCell,
    Text,
    TextStyle,
    VerticalAlign,
)
from ..base import Writer, require
from ..registry import register_writer

__all__ = ["DocxWriter"]

_ALIGNMENT_NAMES = {
    Alignment.LEFT: "LEFT",
    Alignment.CENTER: "CENTER",
    Alignment.RIGHT: "RIGHT",
    Alignment.JUSTIFY: "JUSTIFY",
}

_VALIGN_VALUES = {
    VerticalAlign.TOP: "top",
    VerticalAlign.MIDDLE: "center",
    VerticalAlign.BOTTOM: "bottom",
}

#: Word numbering formats for each universal list style.
_NUM_FMT = {
    ListStyle.BULLET: "bullet",
    ListStyle.ORDERED: "decimal",
    ListStyle.LOWER_ALPHA: "lowerLetter",
    ListStyle.UPPER_ALPHA: "upperLetter",
    ListStyle.LOWER_ROMAN: "lowerRoman",
    ListStyle.UPPER_ROMAN: "upperRoman",
    ListStyle.NONE: "none",
}

_BULLET_GLYPHS = ("\uf0b7", "o", "\uf0a7")


@register_writer
class DocxWriter(Writer):
    """Writes Word documents (``.docx``)."""

    format = "docx"
    extensions = (".docx",)
    aliases = ("word",)
    mime_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    extra = "docx"
    description = "Microsoft Word (Open XML)"

    def write(self, document: Document, stream: BinaryIO, **options: Any) -> None:
        docx = require("docx", "Writing DOCX", extra="docx", package="python-docx")
        builder = _DocxBuilder(docx, document, options)
        native = builder.build()
        native.save(stream)


class _DocxBuilder:
    def __init__(self, docx_module: Any, document: Document, options: Dict[str, Any]) -> None:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        self.docx = docx_module
        self.qn = qn
        self.OxmlElement = OxmlElement
        self.document = document
        self.options = options
        self.template = options.get("template")
        self.embed_images: bool = options.get("embed_images", True)
        self.native = docx_module.Document(self.template) if self.template else docx_module.Document()
        #: Maps a list block's identity to its allocated Word numbering id.
        self._num_ids: Dict[int, int] = {}
        self._next_num_id = 1000
        self._footnote_labels: Dict[str, int] = {}

    # -- entry point ----------------------------------------------------------
    def build(self) -> Any:
        if self.template is None:
            # A fresh python-docx document starts with one empty paragraph.
            self._clear_body()
        self._apply_metadata()
        self._apply_geometry()

        for block in self.document.body:
            self._block(block, self.native)

        if self.document.footnotes:
            self._append_footnote_section()
        return self.native

    def _clear_body(self) -> None:
        body = self.native.element.body
        for child in list(body):
            if child.tag != self.qn("w:sectPr"):
                body.remove(child)

    def _apply_metadata(self) -> None:
        """Copy our metadata across, and clear python-docx's template defaults.

        The default template ships an author of "python-docx" and a 2013 creation date.
        Leaving those in place would attribute the document to the wrong author and
        state a false date, so unspecified fields are cleared rather than inherited.
        """
        import datetime as _datetime

        meta = self.document.metadata
        props = self.native.core_properties

        props.title = meta.title or ""
        props.author = "; ".join(meta.authors) if meta.authors else ""
        props.subject = meta.subject or ""
        props.keywords = ", ".join(meta.keywords) if meta.keywords else ""
        props.comments = meta.description or ""
        props.category = meta.category or ""
        props.last_modified_by = ""
        if meta.language:
            props.language = meta.language

        # A wrong timestamp is worse than a fresh one, so unspecified dates are stamped
        # with now. Pass timestamp=False for byte-reproducible output.
        stamp = self.options.get("timestamp", True)
        now = _datetime.datetime.now(_datetime.timezone.utc).replace(tzinfo=None)
        if meta.created or stamp:
            props.created = meta.created or now
        if meta.modified or stamp:
            props.modified = meta.modified or now
        if not stamp:
            self._clear_timestamps()

    def _clear_timestamps(self) -> None:
        """Remove the date elements outright.

        python-docx raises on ``props.created = None``, so suppressing a date means
        deleting the element; otherwise the template's own 2013 date would remain.
        """
        try:
            element = self.native.core_properties._element
        except AttributeError:  # pragma: no cover
            return
        for tag in ("dcterms:created", "dcterms:modified"):
            try:
                for node in element.findall(self.qn(tag)):
                    element.remove(node)
            except (KeyError, ValueError):  # pragma: no cover
                continue

    def _apply_geometry(self) -> None:
        geometry = self.document.geometry
        if geometry is None:
            return
        from docx.shared import Pt

        for section in self.native.sections:
            section.page_width = Pt(geometry.size.width)
            section.page_height = Pt(geometry.size.height)
            section.left_margin = Pt(geometry.margin_left)
            section.right_margin = Pt(geometry.margin_right)
            section.top_margin = Pt(geometry.margin_top)
            section.bottom_margin = Pt(geometry.margin_bottom)

    # -- blocks ---------------------------------------------------------------
    def _block(self, block: Block, container: Any, list_depth: int = 0) -> None:
        if isinstance(block, Heading):
            self._heading(block, container)
        elif isinstance(block, Paragraph):
            self._paragraph(block, container)
        elif isinstance(block, ListBlock):
            self._list(block, container, list_depth)
        elif isinstance(block, Table):
            self._table(block, container)
        elif isinstance(block, CodeBlock):
            self._code(block, container)
        elif isinstance(block, Quote):
            self._quote(block, container)
        elif isinstance(block, HorizontalRule):
            self._rule(container)
        elif isinstance(block, PageBreak):
            self._page_break(container)
        elif isinstance(block, Image):
            self._image(block, container)
        elif isinstance(block, Section):
            if block.title:
                self._heading(
                    Heading(list(block.title), max(1, min(6, block.level or 1))), container
                )
            for child in block.content:
                self._block(child, container, list_depth)
        elif isinstance(block, Slide):
            if block.title:
                self._heading(Heading.of(block.title, 2), container)
            for child in block.content:
                self._block(child, container, list_depth)
            if block.notes:
                note = self._new_paragraph(container)
                run = note.add_run(f"Notes: {block.notes}")
                run.italic = True
        elif isinstance(block, Page):
            for child in block.content:
                self._block(child, container, list_depth)
        elif isinstance(block, Container):
            if block.role == "sheet" and block.name:
                self._heading(Heading.of(block.name, 2), container)
            for child in block.content:
                self._block(child, container, list_depth)
        elif isinstance(block, Footnote):
            pass  # collected separately
        elif isinstance(block, BlockContainer):
            for child in block.content:
                self._block(child, container, list_depth)
        elif block.text:
            paragraph = self._new_paragraph(container)
            paragraph.add_run(block.text)

    def _new_paragraph(self, container: Any, style: Optional[str] = None) -> Any:
        """Add a paragraph to a document, cell, or header uniformly."""
        if style is not None:
            try:
                return container.add_paragraph(style=style)
            except KeyError:
                # The style is absent from this template; fall back to default.
                pass
        return container.add_paragraph()

    def _heading(self, block: Heading, container: Any) -> None:
        paragraph = self._new_paragraph(container, f"Heading {block.level}")
        self._apply_paragraph_style(paragraph, block.style)
        self._inlines(paragraph, block.content)

    def _paragraph(self, block: Paragraph, container: Any) -> None:
        style_name = block.attrs.get("docx_style") or block.style.style_name
        paragraph = self._new_paragraph(container, style_name if style_name else None)
        self._apply_paragraph_style(paragraph, block.style)
        self._inlines(paragraph, block.content)

    def _code(self, block: CodeBlock, container: Any) -> None:
        from docx.shared import Pt, RGBColor

        for line in block.code.split("\n"):
            paragraph = self._new_paragraph(container, "No Spacing")
            run = paragraph.add_run(line or "")
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1F, 0x23, 0x28)
            fmt = paragraph.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.left_indent = Pt(18)
            self._shade_paragraph(paragraph, "f6f8fa")

    def _quote(self, block: Quote, container: Any) -> None:
        from docx.shared import Pt

        for child in block.content:
            if isinstance(child, Paragraph):
                paragraph = self._new_paragraph(container, "Quote")
                self._apply_paragraph_style(paragraph, child.style)
                paragraph.paragraph_format.left_indent = Pt(36)
                self._inlines(paragraph, child.content)
            else:
                self._block(child, container)
        if block.attribution:
            paragraph = self._new_paragraph(container)
            paragraph.paragraph_format.left_indent = Pt(36)
            run = paragraph.add_run(f"\u2014 {block.attribution}")
            run.italic = True

    def _rule(self, container: Any) -> None:
        """A bottom border on an empty paragraph, which is how Word draws a rule."""
        paragraph = self._new_paragraph(container)
        p_pr = paragraph._p.get_or_add_pPr()
        borders = self.OxmlElement("w:pBdr")
        bottom = self.OxmlElement("w:bottom")
        bottom.set(self.qn("w:val"), "single")
        bottom.set(self.qn("w:sz"), "6")
        bottom.set(self.qn("w:space"), "1")
        bottom.set(self.qn("w:color"), "auto")
        borders.append(bottom)
        p_pr.append(borders)

    def _page_break(self, container: Any) -> None:
        from docx.enum.text import WD_BREAK

        paragraph = self._new_paragraph(container)
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    def _image(self, block: Image, container: Any) -> None:
        from docx.shared import Pt

        paragraph = self._new_paragraph(container)
        if block.style.alignment:
            self._apply_paragraph_style(paragraph, block.style)
        run = paragraph.add_run()
        placed = False
        if self.embed_images and block.data:
            from io import BytesIO

            try:
                width = Pt(block.width) if block.width else None
                run.add_picture(BytesIO(block.data), width=width)
                placed = True
            except Exception:
                placed = False
        if not placed:
            label = block.alt or block.caption or block.src or "image"
            fallback = run.add_text(f"[Image: {label}]")
            del fallback
        if block.caption:
            caption = self._new_paragraph(container, "Caption")
            caption_run = caption.add_run(block.caption)
            caption_run.italic = True

    # -- lists ----------------------------------------------------------------
    def _list(self, block: ListBlock, container: Any, depth: int = 0) -> None:
        num_id = self._numbering_id(block, depth)
        for index, item in enumerate(block.items):
            self._list_item(item, block, container, depth, num_id, index)

    def _list_item(
        self,
        item: ListItem,
        block: ListBlock,
        container: Any,
        depth: int,
        num_id: int,
        index: int,
    ) -> None:
        first = True
        for child in item.content:
            if isinstance(child, ListBlock):
                self._list(child, container, depth + 1)
                continue
            if isinstance(child, Paragraph):
                style = "List Bullet" if not block.ordered else "List Number"
                paragraph = self._new_paragraph(container, style)
                self._apply_paragraph_style(paragraph, child.style)
                self._apply_numbering(paragraph, num_id, depth)
                content: List[Inline] = list(child.content)
                if first and item.checked is not None:
                    marker = "\u2612 " if item.checked else "\u2610 "
                    content = [Text(marker)] + content
                self._inlines(paragraph, content)
                first = False
            else:
                self._block(child, container, depth + 1)

    def _numbering_id(self, block: ListBlock, depth: int) -> int:
        """Allocate (once) a Word numbering definition for a list block."""
        key = id(block)
        existing = self._num_ids.get(key)
        if existing is not None:
            return existing

        numbering = self._numbering_part()
        if numbering is None:
            self._num_ids[key] = 0
            return 0

        abstract_id = self._next_num_id
        num_id = self._next_num_id + 1
        self._next_num_id += 2

        abstract = self.OxmlElement("w:abstractNum")
        abstract.set(self.qn("w:abstractNumId"), str(abstract_id))
        multi = self.OxmlElement("w:multiLevelType")
        multi.set(self.qn("w:val"), "hybridMultilevel")
        abstract.append(multi)

        fmt = _NUM_FMT.get(block.marker_style, "bullet")
        for level in range(9):
            abstract.append(self._level_definition(level, block, fmt))
        numbering.append(abstract)

        num = self.OxmlElement("w:num")
        num.set(self.qn("w:numId"), str(num_id))
        ref = self.OxmlElement("w:abstractNumId")
        ref.set(self.qn("w:val"), str(abstract_id))
        num.append(ref)
        numbering.append(num)

        self._num_ids[key] = num_id
        return num_id

    def _level_definition(self, level: int, block: ListBlock, fmt: str) -> Any:
        lvl = self.OxmlElement("w:lvl")
        lvl.set(self.qn("w:ilvl"), str(level))

        start = self.OxmlElement("w:start")
        start.set(self.qn("w:val"), str(block.start if level == 0 else 1))
        lvl.append(start)

        num_fmt = self.OxmlElement("w:numFmt")
        num_fmt.set(self.qn("w:val"), fmt)
        lvl.append(num_fmt)

        lvl_text = self.OxmlElement("w:lvlText")
        if fmt == "bullet":
            lvl_text.set(self.qn("w:val"), _BULLET_GLYPHS[level % len(_BULLET_GLYPHS)])
        else:
            lvl_text.set(self.qn("w:val"), f"%{level + 1}.")
        lvl.append(lvl_text)

        justify = self.OxmlElement("w:lvlJc")
        justify.set(self.qn("w:val"), "left")
        lvl.append(justify)

        p_pr = self.OxmlElement("w:pPr")
        indent = self.OxmlElement("w:ind")
        indent.set(self.qn("w:left"), str(720 * (level + 1)))
        indent.set(self.qn("w:hanging"), "360")
        p_pr.append(indent)
        lvl.append(p_pr)

        if fmt == "bullet":
            r_pr = self.OxmlElement("w:rPr")
            fonts = self.OxmlElement("w:rFonts")
            font_name = "Symbol" if level % 3 == 0 else ("Courier New" if level % 3 == 1 else "Wingdings")
            fonts.set(self.qn("w:ascii"), font_name)
            fonts.set(self.qn("w:hAnsi"), font_name)
            fonts.set(self.qn("w:hint"), "default")
            r_pr.append(fonts)
            lvl.append(r_pr)
        return lvl

    def _numbering_part(self) -> Any:
        try:
            return self.native.part.numbering_part.element
        except (AttributeError, KeyError, NotImplementedError):
            return None

    def _apply_numbering(self, paragraph: Any, num_id: int, level: int) -> None:
        if not num_id:
            return
        p_pr = paragraph._p.get_or_add_pPr()
        # Replace any numbering the style already applied.
        for existing in p_pr.findall(self.qn("w:numPr")):
            p_pr.remove(existing)
        num_pr = self.OxmlElement("w:numPr")
        ilvl = self.OxmlElement("w:ilvl")
        ilvl.set(self.qn("w:val"), str(min(8, level)))
        num_pr.append(ilvl)
        num_ref = self.OxmlElement("w:numId")
        num_ref.set(self.qn("w:val"), str(num_id))
        num_pr.append(num_ref)
        p_pr.append(num_pr)

    # -- tables ---------------------------------------------------------------
    def _table(self, block: Table, container: Any) -> None:
        rows, columns = block.dimensions
        if not rows or not columns:
            return

        style = block.attrs.get("docx_style") or self.options.get("table_style", "Table Grid")
        try:
            native = container.add_table(rows=rows, cols=columns, style=style)
        except (KeyError, ValueError):
            native = container.add_table(rows=rows, cols=columns)

        if block.column_widths:
            self._apply_column_widths(native, block.column_widths)

        for row_index, row in enumerate(block.rows):
            native_row = native.rows[row_index]
            if row.is_header:
                self._mark_header_row(native_row)
            column = 0
            for cell in row.cells:
                if column >= columns:
                    break
                target = native_row.cells[column]
                span_end = min(columns - 1, column + cell.colspan - 1)
                if span_end > column:
                    target = target.merge(native_row.cells[span_end])
                if cell.rowspan > 1:
                    last_row = min(rows - 1, row_index + cell.rowspan - 1)
                    if last_row > row_index:
                        try:
                            target = target.merge(native.rows[last_row].cells[column])
                        except (IndexError, ValueError):
                            pass
                self._fill_cell(target, cell, header=row.is_header)
                column = span_end + 1

        if block.caption:
            caption = self._new_paragraph(container, "Caption")
            run = caption.add_run(block.caption)
            run.italic = True

    def _fill_cell(self, native_cell: Any, cell: TableCell, header: bool = False) -> None:
        # A new python-docx cell already contains one empty paragraph; reuse it.
        existing = native_cell.paragraphs[0]
        first = True
        for child in cell.content:
            if isinstance(child, Paragraph):
                paragraph = existing if first else native_cell.add_paragraph()
                self._apply_paragraph_style(paragraph, child.style)
                content = child.content
                if header:
                    content = [
                        Text(node.text, node.style.merge(TextStyle(bold=True)))
                        if isinstance(node, Text)
                        else node
                        for node in content
                    ]
                self._inlines(paragraph, content)
                first = False
            else:
                self._block(child, native_cell)

        if cell.valign is not None:
            value = _VALIGN_VALUES.get(cell.valign)
            if value:
                tc_pr = native_cell._tc.get_or_add_tcPr()
                v_align = self.OxmlElement("w:vAlign")
                v_align.set(self.qn("w:val"), value)
                tc_pr.append(v_align)
        if cell.background:
            self._shade_cell(native_cell, cell.background.lstrip("#"))

    def _mark_header_row(self, native_row: Any) -> None:
        tr_pr = native_row._tr.get_or_add_trPr()
        header = self.OxmlElement("w:tblHeader")
        header.set(self.qn("w:val"), "true")
        tr_pr.append(header)

    def _apply_column_widths(self, native: Any, widths: Sequence[float]) -> None:
        from docx.shared import Pt

        total = sum(w for w in widths if w > 0)
        if total <= 0:
            return
        native.autofit = False
        for index, width in enumerate(widths):
            if width <= 0:
                continue
            try:
                for row in native.rows:
                    row.cells[index].width = Pt(width)
            except IndexError:
                break

    def _shade_cell(self, native_cell: Any, fill: str) -> None:
        tc_pr = native_cell._tc.get_or_add_tcPr()
        shade = self.OxmlElement("w:shd")
        shade.set(self.qn("w:val"), "clear")
        shade.set(self.qn("w:color"), "auto")
        shade.set(self.qn("w:fill"), fill)
        tc_pr.append(shade)

    def _shade_paragraph(self, paragraph: Any, fill: str) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        shade = self.OxmlElement("w:shd")
        shade.set(self.qn("w:val"), "clear")
        shade.set(self.qn("w:color"), "auto")
        shade.set(self.qn("w:fill"), fill)
        p_pr.append(shade)

    # -- inline ---------------------------------------------------------------
    def _inlines(self, paragraph: Any, content: Sequence[Inline]) -> None:
        for node in content:
            self._inline(paragraph, node)

    def _inline(self, paragraph: Any, node: Inline) -> None:
        if isinstance(node, Text):
            self._text_run(paragraph, node.text, node.style)
        elif isinstance(node, Link):
            self._hyperlink(paragraph, node)
        elif isinstance(node, LineBreak):
            paragraph.add_run().add_break()
        elif isinstance(node, InlineImage):
            self._inline_image(paragraph, node)
        elif isinstance(node, Math):
            self._text_run(paragraph, node.latex, TextStyle(italic=True))
        elif isinstance(node, FootnoteRef):
            self._text_run(
                paragraph,
                node.label or f"[{node.identifier}]",
                TextStyle(superscript=True),
            )
        elif isinstance(node, DynamicField):
            self._field(paragraph, node)
        elif node.text:
            paragraph.add_run(node.text)

    def _text_run(self, paragraph: Any, text: str, style: TextStyle) -> Any:
        from docx.shared import Pt, RGBColor

        run = paragraph.add_run(text)
        font = run.font
        if style.bold is not None:
            run.bold = style.bold
        if style.italic is not None:
            run.italic = style.italic
        if style.underline is not None:
            run.underline = style.underline
        if style.strike is not None:
            font.strike = style.strike
        if style.small_caps is not None:
            font.small_caps = style.small_caps
        if style.superscript:
            font.superscript = True
        if style.subscript:
            font.subscript = True
        if style.is_monospace:
            font.name = style.font_family or "Consolas"
            if style.font_size is None:
                font.size = Pt(9.5)
        elif style.font_family:
            font.name = style.font_family
        if style.font_size:
            font.size = Pt(style.font_size)
        if style.color:
            try:
                red, green, blue = _hex_rgb(style.color)
                font.color.rgb = RGBColor(red, green, blue)
            except ValueError:
                pass
        if style.highlight:
            from docx.enum.text import WD_COLOR_INDEX

            font.highlight_color = WD_COLOR_INDEX.YELLOW
        return run

    def _hyperlink(self, paragraph: Any, link: Link) -> None:
        """Build a ``w:hyperlink`` element, which python-docx cannot add directly."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        href = link.href or ""
        element = self.OxmlElement("w:hyperlink")
        if href.startswith("#"):
            element.set(self.qn("w:anchor"), href[1:])
        elif href:
            rel_id = paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True)
            element.set(self.qn("r:id"), rel_id)

        # Build the runs inside a throwaway paragraph, then move them across.
        start = len(paragraph.runs)
        for node in link.content:
            self._inline(paragraph, node)
        moved = paragraph.runs[start:]
        if not moved:
            run = self._text_run(paragraph, href, TextStyle())
            moved = [run]
        for run in moved:
            run.font.underline = True
            self._apply_hyperlink_colour(run)
            paragraph._p.remove(run._r)
            element.append(run._r)
        paragraph._p.append(element)

    def _apply_hyperlink_colour(self, run: Any) -> None:
        from docx.shared import RGBColor

        if run.font.color is None or run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)

    def _inline_image(self, paragraph: Any, node: InlineImage) -> None:
        from docx.shared import Pt

        if self.embed_images and node.data:
            from io import BytesIO

            try:
                run = paragraph.add_run()
                run.add_picture(
                    BytesIO(node.data), width=Pt(node.width) if node.width else None
                )
                return
            except Exception:
                pass
        if node.alt:
            self._text_run(paragraph, f"[{node.alt}]", TextStyle(italic=True))

    def _field(self, paragraph: Any, node: DynamicField) -> None:
        """Emit a live Word field so page numbers stay correct after editing."""
        instructions = {
            "page-number": "PAGE",
            "page-count": "NUMPAGES",
            "date": "DATE",
            "time": "TIME",
            "title": "TITLE",
            "author": "AUTHOR",
            "filename": "FILENAME",
        }
        instruction = instructions.get(node.kind)
        if instruction is None:
            if node.fallback:
                self._text_run(paragraph, node.fallback, node.style)
            return

        run = paragraph.add_run()
        begin = self.OxmlElement("w:fldChar")
        begin.set(self.qn("w:fldCharType"), "begin")
        instr = self.OxmlElement("w:instrText")
        instr.set(self.qn("xml:space"), "preserve")
        instr.text = f" {instruction} "
        end = self.OxmlElement("w:fldChar")
        end.set(self.qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)

    # -- footnotes ------------------------------------------------------------
    def _append_footnote_section(self) -> None:
        """Render footnotes as an endnote-style section.

        Word's real footnote store requires editing ``footnotes.xml`` and its content
        types; a labelled section at the end keeps the text and the references without
        risking a corrupt package.
        """
        from docx.shared import Pt

        self._rule(self.native)
        for note in self.document.footnotes:
            for index, child in enumerate(note.content):
                if isinstance(child, Paragraph):
                    paragraph = self._new_paragraph(self.native)
                    paragraph.paragraph_format.space_after = Pt(2)
                    if index == 0:
                        marker = self._text_run(
                            paragraph, note.identifier, TextStyle(superscript=True)
                        )
                        del marker
                        paragraph.add_run(" ")
                    self._inlines(paragraph, child.content)
                else:
                    self._block(child, self.native)

    def _apply_paragraph_style(self, paragraph: Any, style: ParagraphStyle) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        fmt = paragraph.paragraph_format
        if style.alignment is not None:
            name = _ALIGNMENT_NAMES.get(style.alignment)
            if name is not None:
                fmt.alignment = getattr(WD_ALIGN_PARAGRAPH, name)
        if style.space_before is not None:
            fmt.space_before = Pt(style.space_before)
        if style.space_after is not None:
            fmt.space_after = Pt(style.space_after)
        if style.line_spacing is not None:
            fmt.line_spacing = (
                style.line_spacing if style.line_spacing <= 4 else Pt(style.line_spacing)
            )
        if style.indent_left is not None:
            fmt.left_indent = Pt(style.indent_left)
        if style.indent_right is not None:
            fmt.right_indent = Pt(style.indent_right)
        if style.first_line_indent is not None:
            fmt.first_line_indent = Pt(style.first_line_indent)
        if style.background:
            self._shade_paragraph(paragraph, style.background.lstrip("#"))


def _hex_rgb(color: str) -> Tuple[int, int, int]:
    value = color.lstrip("#")
    if len(value) != 6:
        raise ValueError(f"Not a 6-digit hex colour: {color!r}")
    return (int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))
